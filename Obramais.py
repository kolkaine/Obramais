# app.py — Orçamento de Materiais de Obra (m² / m³) com exportação Excel, PDF, CSV e DOCX
#
# Execute localmente:  streamlit run app.py
# Dependências sugeridas (requirements.txt):
#   streamlit
#   pandas
#   reportlab
#   openpyxl
#   XlsxWriter
#   python-docx
#
# Observação:
# - Inclui parâmetros típicos de cobertura por material como referência.
# - Permite inserir custos adicionais (mão de obra, impostos).
# - Exporta para Excel, PDF, CSV e DOCX.

import streamlit as st
import pandas as pd
import io
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# ---------- Funções utilitárias ----------
def calc_qtd_necessaria(medida: float, cobertura_por_unidade: float, desperdicio_pct: float) -> float:
    if cobertura_por_unidade <= 0:
        return 0.0
    fator = 1 + (desperdicio_pct or 0) / 100.0
    return (medida / cobertura_por_unidade) * fator

def add_item(material: str, medida: float, cobertura: float, unidade: str,
             desperdicio: float, preco_unit: float, especificacoes: dict):
    qtd = calc_qtd_necessaria(medida, cobertura, desperdicio)
    qtd = round(qtd, 3)
    subtotal = round((preco_unit or 0.0) * qtd, 2)
    item = {
        "Material": material,
        "Medida": medida,                       # m² ou m³ conforme o caso
        "Cobertura por Unidade": cobertura,     # quanto 1 unidade cobre (m² ou m³)
        "Unidade": unidade,                     # lata, saco, peça, m³, m², un etc.
        "Desperdício (%)": desperdicio,
        "Qtd Necessária": qtd,
        "Preço Unitário": float(preco_unit or 0.0),
        "Subtotal": subtotal,
    }
    item.update(especificacoes)  # adiciona campos extras específicos do material
    st.session_state["items"].append(item)

def df_resumo() -> pd.DataFrame:
    data = st.session_state.get("items", [])
    return pd.DataFrame(data) if data else pd.DataFrame()

def make_excel_bytes(df: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Orçamento")
    return buffer.getvalue()

def make_pdf_bytes(df: pd.DataFrame, projeto: str, cliente: str, responsavel: str, custos_extra: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    styles = getSampleStyleSheet()
    story = []

    titulo = f"Orçamento de Materiais — {projeto}"
    story.append(Paragraph(titulo, styles["Title"]))
    meta = f"Cliente: {cliente or '-'} | Responsável: {responsavel or '-'} | Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    story.append(Paragraph(meta, styles["Normal"]))
    story.append(Spacer(1, 12))

    if not df.empty:
        header = list(df.columns)
        data = [header] + [[str(r[c]) for c in df.columns] for _, r in df.iterrows()]
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))
        total = float(df["Subtotal"].sum()) if "Subtotal" in df.columns else 0.0
        total_final = total + float(custos_extra.get("mao_obra", 0)) + float(custos_extra.get("impostos", 0))
        story.append(Paragraph(f"<b>Materiais:</b> R$ {total:.2f}", styles["Normal"]))
        story.append(Paragraph(f"<b>Mão de obra:</b> R$ {float(custos_extra.get('mao_obra',0)):.2f}", styles["Normal"]))
        story.append(Paragraph(f"<b>Impostos:</b> R$ {float(custos_extra.get('impostos',0)):.2f}", styles["Normal"]))
        story.append(Paragraph(f"<b>Total:</b> R$ {total_final:.2f}", styles["Heading3"]))

    doc.build(story)
    return buffer.getvalue()

def make_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False, sep=";").encode("utf-8")

def make_docx_bytes(df: pd.DataFrame, projeto: str, cliente: str, responsavel: str, custos_extra: dict) -> bytes:
    doc = Document()
    doc.add_heading(f"Orçamento de Materiais — {projeto}", 0)
    doc.add_paragraph(f"Cliente: {cliente or '-'}")
    doc.add_paragraph(f"Responsável: {responsavel or '-'}")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    if not df.empty:
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        total = float(df["Subtotal"].sum()) if "Subtotal" in df.columns else 0.0
        total_final = total + float(custos_extra.get("mao_obra", 0)) + float(custos_extra.get("impostos", 0))
        doc.add_paragraph(f"Materiais: R$ {total:.2f}")
        doc.add_paragraph(f"Mão de obra: R$ {float(custos_extra.get('mao_obra',0)):.2f}")
        doc.add_paragraph(f"Impostos: R$ {float(custos_extra.get('impostos',0)):.2f}")
        doc.add_heading(f"Total: R$ {total_final:.2f}", level=1)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ---------- App ----------
st.title("📐 Orçamento de Materiais de Construção")

# Inicializa session_state["items"]
if "items" not in st.session_state:
    st.session_state["items"] = []

projeto = st.text_input("Nome do Projeto")
cliente = st.text_input("Cliente")
responsavel = st.text_input("Responsável")

st.subheader("Adicionar Material")

material = st.selectbox("Material", [
    "Tinta", "Cimento", "Ladrilhos", "Madeira", "Pregos", "Cola",
    "Canos", "Janelas", "Gesso", "Blocos", "Tijolos", "Areia", "Brita"
])

medida = st.number_input("Área/Volume (m² ou m³)", min_value=0.0, step=0.1)
cobertura = st.number_input("Cobertura por unidade", min_value=0.0, step=0.1)
unidade = st.text_input("Unidade de Medida (ex: lata, saco, m², m³, un)")
desperdicio = st.number_input("Desperdício (%)", min_value=0.0, value=5.0, step=1.0)
preco_unit = st.number_input("Preço unitário (R$)", min_value=0.0, step=0.1)

# Campos extras por material
especificacoes = {}
if material == "Madeira":
    especificacoes["Medidas (m)"] = st.text_input("Medidas da madeira (ex: 2.5 x 0.1 x 0.03)")
elif material == "Cola":
    especificacoes["Tipo"] = st.selectbox("Tipo de cola", ["Hidráulica", "Madeira", "Universal", "PVC"])
elif material == "Canos":
    especificacoes["Diâmetro"] = st.text_input("Diâmetro do cano (mm)")
    especificacoes["Uso"] = st.selectbox("Uso do cano", ["Pia", "Privada", "Esgoto", "Água fria", "Água quente", "Pluvial"])
elif material == "Janelas":
    especificacoes["Altura (m)"] = st.number_input("Altura da janela (m)", min_value=0.0, step=0.1)
    especificacoes["Largura (m)"] = st.number_input("Largura da janela (m)", min_value=0.0, step=0.1)
elif material == "Blocos":
    especificacoes["Medida do bloco (cm)"] = st.text_input("Medida do bloco (ex: 14x19x39)")
elif material == "Tijolos":
    especificacoes["Medida do tijolo (cm)"] = st.text_input("Medida do tijolo (ex: 9x19x29)")
elif material == "Areia":
    especificacoes["Tipo de areia"] = st.selectbox("Tipo de areia", ["Fina", "Média", "Grossa"])
elif material == "Brita":
    especificacoes["Tipo de brita"] = st.selectbox("Tipo de brita", ["Nº 0", "Nº 1", "Nº 2", "Nº 3"])

if st.button("Adicionar Material"):
    add_item(material, medida, cobertura, unidade, desperdicio, preco_unit, especificacoes)
    st.success(f"{material} adicionado ao orçamento!")

st.subheader("Resumo dos Itens")
df = df_resumo()
if not df.empty:
    st.dataframe(df)

    st.subheader("Custos Extras")
    mao_obra = st.number_input("Custo de mão de obra (R$)", min_value=0.0, step=0.1)
    impostos = st.number_input("Impostos (R$)", min_value=0.0, step=0.1)
    custos_extra = {"mao_obra": mao_obra, "impostos": impostos}

    total = float(df["Subtotal"].sum()) if "Subtotal" in df.columns else 0.0
    total_final = total + mao_obra + impostos
    st.write(f"**Total Materiais:** R$ {total:.2f}")
    st.write(f"**Total Final (com mão de obra e impostos):** R$ {total_final:.2f}")

    st.subheader("Exportar Orçamento")
    st.download_button("📊 Baixar Excel", make_excel_bytes(df), "orcamento.xlsx")
    st.download_button("📄 Baixar PDF", make_pdf_bytes(df, projeto, cliente, responsavel, custos_extra), "orcamento.pdf")
    st.download_button("📝 Baixar CSV", make_csv_bytes(df), "orcamento.csv")
    st.download_button("📑 Baixar DOCX", make_docx_bytes(df, projeto, cliente, responsavel, custos_extra), "orcamento.docx")
else:
    st.info("Nenhum item adicionado ainda.")